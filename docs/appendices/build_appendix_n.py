from pathlib import Path
import json, math
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parent
# Columns: System, Case Manager, Agency Focal Person, Administrator, OFW.
# A is automated processing, not a human permission. ! is an implementation discrepancy.
modules = []
def module(name, source, *rows):
    modules.append({'name': name, 'source': source, 'rows': rows})

module('Authentication', 'routes/auth.php; app/Http/Requests/Auth/LoginRequest.php',
 ('Log in with an existing account', '1234'), ('Log out', '1234'),
 ('Request password reset and reset password', '1234'),
 ('Register staff account through an invitation', '123'),
 ('Register OFW account after verified intake or tracking', '4'),
 ('Verify account email / resend verification link', '1234'))
module('Authorization', 'app/Http/Middleware/CheckRole.php; bootstrap/app.php; app/Http/Controllers/CaseController.php::authorizeCaseAccess',
 ('Validate user role and enforce route access', '0'),
 ('Check record ownership, agency scope and verified public access', '0'),
 ('Check account activity and required MFA', '0'))
module('Email OTP', 'app/Http/Controllers/TrackController.php; app/Http/Controllers/IntakeController.php; app/Http/Controllers/EmailChangeController.php',
 ('Generate, send and validate email OTP', '0'),
 ('Request / resend tracking or intake OTP [1]', '1234'),
 ('Submit tracking or intake OTP for verification [1]', '1234'),
 ('Request and verify an account email change', '1234'))
module('Multi-Factor Authentication', 'app/Http/Controllers/MfaController.php; config/mfa.php; routes/auth.php',
 ('Set up authenticator and enable MFA [2]', '1234'),
 ('Complete required login MFA / recovery challenge [2]', '123'),
 ('Regenerate recovery codes / disable MFA [2]', '1234'))
module('Profile', 'app/Http/Controllers/ProfileController.php; app/Http/Controllers/OfwProfileController.php; routes/auth.php',
 ('View own profile', '1234'), ('Update own permitted profile fields [3]', '1234'),
 ('Change own password', '1234'),
 ('Manage avatar and notification preferences through shared profile [2]', '1234'),
 ('Delete own account through shared profile endpoint [2]', '1234'))
module('Landing Page and Public Information', 'routes/web.php; resources/js/Pages/Welcome.jsx',
 ('View system overview', '1234'), ('View public agencies and services', '1234'),
 ('View frequently asked questions', '1234'), ('Submit contact form', '1234'),
 ('View privacy policy and terms', '1234'))
module('Public Intake', 'app/Http/Controllers/IntakeController.php; app/Services/IntakeService.php; routes/web.php',
 ('Submit a self-filed assistance request [1]', '1234'),
 ('Check returning client / duplicate intake after verification [1]', '1234'),
 ('Prefill own client details when signed in', '4'),
 ('Create intake draft and send acknowledgement', '0'))
module('Tracking Input', 'app/Http/Controllers/TrackController.php',
 ('Enter tracking number and matching email [1]', '1234'),
 ('Submit tracking request and verify access [1]', '1234'))
module('Tracking View and Timeline', 'app/Services/TrackingService.php; app/Http/Controllers/OfwDashboardController.php; app/Services/ReferralService.php::getReferralTimeline',
 ('View case status and overview [1, 4]', '1234'),
 ('View client-facing case history and timeline [1, 4]', '1234'),
 ('View referred agencies and referral status [1, 4]', '1234'),
 ('View agency milestones and requirements [1, 4]', '1234'),
 ('View internal referral timeline [4]', '123'))
module('Client', 'app/Http/Controllers/ClientController.php; app/Services/CaseService.php::publishDraft; app/Services/CaseService.php::updateDraft',
 ('Create client through case entry / publication', '13'),
 ('View, search and filter client records [4]', '123'),
 ('Update linked address, employment and next of kin through case draft [5]', '13'),
 ('Upload or remove client profile picture [4]', '123'),
 ('Export permitted client records to Excel [4]', '123'),
 ('Access client selection API [10]', '13!?'))
module('Case', 'app/Http/Controllers/CaseController.php; app/Services/CaseService.php; routes/web.php',
 ('Create case / save initial draft', '13'),
 ('View, search and filter staff case list', '13'),
 ('View internal case details [4]', '123'),
 ('Update case fields', '13'),
 ('View / edit / save eligible draft [5]', '13'),
 ('Publish a complete draft', '13'),
 ('Delete eligible draft [5]', '13'),
 ('Export case record to PDF / case list to Excel', '13'))
module('Intake Review', 'app/Http/Controllers/CaseController.php::intakeQueue; app/Http/Controllers/CaseController.php::reviewIntake; app/Services/CaseService.php',
 ('View pending intake queue', '13'), ('Review and correct intake draft', '13'),
 ('Accept intake by publishing case', '13'), ('Reject intake with reason', '13'))
module('Case Status and Archive', 'app/Services/CaseService.php::canClose; app/Services/CaseService.php::toggleCaseStatus; app/Services/CaseService.php::deleteArchivedCase',
 ('Update case status / close eligible case [6]', '13'),
 ('Reopen closed case', '13'), ('Archive closed case / unarchive case', '13'),
 ('Soft-delete archived case with reason', '13'),
 ('View trash and restore deleted case', '13'),
 ('Purge expired trash under retention schedule', '0'))
module('Case Document', 'app/Http/Controllers/CaseDocumentController.php; routes/web.php',
 ('Upload case document', '13'), ('View / download permitted document [4]', '123'),
 ('Soft-delete case document', '13'))
module('Referral', 'app/Http/Controllers/ReferralController.php; app/Http/Requests/StoreReferralRequest.php; app/Services/ReferralService.php::updateStatus',
 ('Create referral', '13'), ('View, search and filter referral list [4, 8]', '123!'),
 ('View internal referral details [4]', '123'),
 ('Accept referral as receiving agency', '2'), ('Reject referral as receiving agency', '2'),
 ('Update permitted referral status [6]', '123'),
 ('Export referrals to Excel [4, 8]', '123!'))
module('Milestone', 'app/Http/Controllers/ReferralController.php::addMilestone; app/Services/ReferralService.php::addMilestone; app/Services/TrackingService.php',
 ('Create milestone on eligible referral [6]', '123'),
 ('View milestones [1, 4]', '1234'))
module('Comment', 'app/Http/Controllers/ReferralController.php::addComment; app/Http/Controllers/ReferralController.php::replyToComment',
 ('Create internal referral comment [4]', '123'),
 ('View internal referral comments [4]', '123'),
 ('Reply to internal referral comment [4]', '123'))
module('Referral Attachment', 'app/Http/Controllers/ReferralController.php::addAttachment; app/Http/Controllers/ReferralController.php::deleteAttachment',
 ('Upload attachment to accessible referral [4]', '123'),
 ('View / download referral attachment [4]', '123'),
 ('Replace attachment and view version history [4]', '123'),
 ('Soft-delete own uploaded attachment [4]', '123'))
module('Referral Services and Requirements', 'app/Http/Controllers/ReferralController.php::addService; app/Http/Controllers/ReferralController.php::updateRequirement',
 ('Add / remove receiving agency service on referral', '2'),
 ('Add / update / delete referral requirements', '2'),
 ('View referral service details [4]', '123'))
module('Client Requests and Secure Inbox', 'app/Http/Controllers/ReferralClientRequestController.php; app/Services/ReferralClientRequestService.php',
 ('Create document / question / information request', '2'),
 ('View request history and messages [4]', '123'),
 ('Send agency message on own referral', '2'),
 ('Complete / cancel / reopen eligible request', '2'),
 ('Issue or reissue secure client access link', '2'),
 ('Revoke client access link [4]', '123'),
 ('Read and reply through verified client link [1]', '1234'),
 ('Upload / replace requested client attachment [1]', '1234'),
 ('Download permitted request attachments [1, 4]', '1234'))
module('Activity and Audit Log', 'app/Http/Controllers/AuditLogController.php; app/Observers/AuditObserver.php; routes/auth.php',
 ('Automatically record auditable activity', '0'),
 ('View and filter activity / audit records [7]', '123'),
 ('View case-level activity history', '13'),
 ('View referral-level activity history [4]', '123'),
 ('Export audit logs with date range', '3'))
module('Notification', 'app/Services/NotificationService.php; app/Services/ReferralService.php; app/Services/CaseService.php; app/Http/Controllers/OfwDashboardController.php',
 ('Send referral / status / milestone notifications', '0'),
 ('Send tracking number / intake publication email', '0'),
 ('View own notifications and mark as read', '1234'),
 ('Mark all staff notifications as read', '123'))
module('Dashboard', 'app/Http/Controllers/DashboardController.php; app/Http/Controllers/OfwDashboardController.php',
 ('View staff dashboard and case / referral summary [4]', '123'),
 ('View own OFW dashboard and case summary', '4'),
 ('Open own client-facing case details', '4'))
module('Reporting', 'app/Http/Controllers/ReportsController.php; app/Services/ReportsService.php; app/Services/Reports/ReportsExportService.php',
 ('View report dashboard / generate on-screen report [8]', '123!'),
 ('Filter reports by date range / date scope [8]', '123!'),
 ('Select agency filter (agency users are fixed to own agency)', '13'),
 ('Filter reports by province / city [8]', '123!'),
 ('View case and referral breakdowns [4, 8]', '123!'),
 ('View client demographics and descriptive statistics [4, 8]', '123!'),
 ('Display line graphs and bar charts [4, 8]', '123!'),
 ('Export permitted reports to PDF / Excel [4]', '123'))
module('Overdue Referral', 'app/Http/Controllers/Admin/OverdueReferralController.php; routes/console.php',
 ('View and filter overdue referrals [4]', '123'),
 ('Request reminder delivery', '13'),
 ('Deliver requested overdue notifications', '0'))
module('User Administration', 'app/Http/Controllers/AdminUserController.php; routes/web.php',
 ('Create / view / search user accounts', '3'),
 ('Update user details, role and account status', '3'),
 ('Soft-delete / reactivate user', '3'),
 ('Invite staff / resend or cancel invitation', '3'),
 ('Verify account / reset user MFA', '3'),
 ('Change user email through OTP verification', '3'))
module('Agency Administration', 'app/Http/Controllers/AdminAgencyController.php; routes/web.php',
 ('Create agency', '3'), ('View agency administration list', '3'),
 ('View agency administrative detail [9]', '13'),
 ('Update agency', '3'), ('Deactivate / reactivate agency [9]', '3'))
module('Service Catalogue', 'app/Http/Controllers/AgencyServiceController.php; app/Http/Controllers/AdminServiceController.php',
 ('Create service and requirements [4]', '23'),
 ('View / search service catalogue [4]', '23'),
 ('Update service and requirements [4]', '23'),
 ('Soft-delete service', '23'))
module('Case Reference Data', 'app/Http/Controllers/Admin/AdminCaseCategoryController.php; app/Http/Controllers/Admin/AdminCaseStatusController.php; app/Http/Controllers/Admin/AdminCaseIssueController.php',
 ('Manage case categories / reactivate category', '3'),
 ('Manage case status definitions', '3'),
 ('Manage case issues / reactivate issue', '3'),
 ('Quick-create case issue during case entry', '13'))
module('Stakeholder Directory', 'app/Http/Controllers/StakeholderController.php; routes/web.php',
 ('View stakeholder directory', '13'), ('View stakeholder details', '13'))
module('Survey', 'app/Http/Controllers/SurveyFormController.php; app/Http/Controllers/SurveyResponseController.php; app/Http/Controllers/PublicSurveyController.php',
 ('Create / view / update own agency survey forms', '2'),
 ('Activate / delete own agency survey form', '2'),
 ('View survey response statistics and details [4]', '123'),
 ('View / submit survey using valid invitation [1]', '1234'),
 ('Send survey invitation after eligible completion', '0'))
module('System Settings and Operations', 'app/Http/Controllers/SystemSettingsController.php; app/Http/Controllers/Admin; routes/web.php',
 ('View / update system settings', '3'),
 ('Reindex AI assistant knowledge', '3'),
 ('View / download application logs', '3'),
 ('View / toggle maintenance mode', '3'),
 ('View / update security settings', '3'),
 ('View active sessions / terminate session', '3'),
 ('View email delivery logs / resend eligible email', '3'))
module('Administrative Data Export', 'app/Http/Controllers/Admin/DataExportController.php; routes/web.php',
 ('View data export options', '3'), ('Export full administrative workbook', '3'))
module('Helpdesk and Onboarding', 'routes/web.php; app/Http/Controllers/OnboardingController.php; app/Services/OnboardingService.php',
 ('Browse / search / read helpdesk articles', '1234'),
 ('View / complete / skip / replay onboarding state [2]', '1234'),
 ('Update onboarding steps, guides and checklist [2]', '1234'))
module('AI Assistant', 'app/Http/Controllers/ChatbotController.php; app/Providers/AppServiceProvider.php::boot',
 ('Ask for tracking guidance (no live case lookup)', '1234'),
 ('Ask for service, agency and system guidance', '1234'),
 ('Apply topic / prompt-injection guards and audience scope', '0'),
 ('Enforce request rate and response length limits', '0'))

notes = [
 'Public and token-based features are not restricted to the OFW role. An asterisk in a staff column does not bypass matching-email OTP, a valid survey invitation, or a valid client-request link. The OFW column includes public clients and registered OFW users; account-only actions require registration.',
 'Shared profile, MFA setup and onboarding endpoints accept authenticated verified users, including OFWs, although the OFW portal has its own profile. Login MFA is required for configured roles (default: Administrator, Case Manager and Agency Focal Person); OFW enrolment is not required by default. Endpoint access does not imply an OFW-specific screen, onboarding tour or default OFW login challenge.',
 'OFW profile editing covers own contact number, address, employment, next of kin and password. Identity fields in the OFW profile are read-only. Staff use the shared profile editor.',
 'Case Managers and Administrators generally have system-wide operational access. Agency Focal Persons are restricted to their agency referrals and associated records; the internal case-detail route requires an active referral, and case documents require a referral to their agency. OFWs use their own client-facing cases or verified public tracking. Report panels differ by role.',
 'Client creation is part of case publication or public intake, not a standalone client-create page. Client updates listed here are limited to implemented workflows; no general client update/delete endpoint exists. Internal draft editing is ownership-restricted; self-filed drafts are reviewable by Case Managers and Administrators. Draft deletion is separate from archived-case trash.',
 'Case closure is blocked while any referral is neither completed nor rejected. Only the receiving agency accepts or rejects a referral. Case Managers and Administrators can make other permitted transitions after agency engagement. Milestones cannot be added to completed referrals; agency milestone entry requires processing or for-compliance status. No direct milestone edit/delete or comment edit/delete function is listed because no such route is registered.',
 'Activity Log and Audit Logs share the same implementation. Current controller code gives Case Managers and Administrators the unscoped log query; Agency Focal Persons get an agency-scoped query. Export is Administrator-only. This differs from comments that describe narrower Case Manager access.',
 'The exclamation mark identifies an access discrepancy, not an approved OFW permission. The reports page accepts authenticated verified OFWs and defaults to the Case Manager payload without client ownership scope. Report export rejects OFWs. Referral list, creation-page and referral Excel export routes also lack an explicit OFW denial; their list/export queries restrict agencies but do not restrict OFWs. Referral creation submission still requires Case Manager or Administrator. Resolve these discrepancies before treating this matrix as an approved security policy.',
 'Administrative routes also apply the configured IP whitelist, including the Case Manager agency-detail route. Agency removal sets inactive and deleted flags; the default agency is protected. Service removal uses the model soft-delete behavior.',
 'Client selection API routes require authentication and verification but do not apply the staff role gate or agency/client scope in ClientSelectController. Exclamation marks for Agency Focal Person and OFW flag this discrepancy; ordinary client pages retain their separate route and controller checks.',
]

def xml(tag, **attrs):
    el = OxmlElement('w:'+tag)
    for k,v in attrs.items(): el.set(qn('w:'+k), str(v))
    return el

def build():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.top_margin, sec.bottom_margin = Inches(.55), Inches(.55)
    sec.left_margin = sec.right_margin = Inches(.625)
    for style in doc.styles:
        if style.type in (1,2):
            style.font.name = 'Times New Roman'
            style.font.color.rgb = RGBColor(0,0,0)
    normal = doc.styles['Normal']
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1
    title = doc.add_paragraph(style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    r = title.add_run('Appendix N\nList of Modules')
    r.bold = True; r.font.size = Pt(12)
    p = doc.add_paragraph('Legend: * = implemented access; blank = no listed access. System * = automated processing.\n! = access discrepancy (note 8). Bracketed numbers refer to the notes after the table.')
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs: r.font.size = Pt(9)
    widths = [.95,2.65,.55,.80,.85,.85,.60]
    table = doc.add_table(rows=1, cols=7)
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c,w in zip(table.columns,widths): c.width=Inches(w)
    borders=xml('tblBorders')
    for edge in ['top','left','bottom','right','insideH','insideV']:
        borders.append(xml(edge,val='single',sz='4',color='000000'))
    table._tbl.tblPr.append(borders)
    table._tbl.tblPr.append(xml('tblLayout',type='fixed'))
    margins=xml('tblCellMar')
    for edge,val in [('top',35),('bottom',35),('left',60),('right',60)]:
        margins.append(xml(edge,w=val,type='dxa'))
    table._tbl.tblPr.append(margins)
    def setrow(row, texts, bold=False, keep=False):
        row._tr.get_or_add_trPr().append(xml('cantSplit'))
        for i,(c,t) in enumerate(zip(row.cells,texts)):
            c.width=Inches(widths[i]); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=c.paragraphs[0]; p.paragraph_format.keep_with_next=keep
            p.alignment=WD_ALIGN_PARAGRAPH.LEFT if i<2 else WD_ALIGN_PARAGRAPH.CENTER
            r=p.add_run(t); r.bold=bold; r.font.size=Pt(10 if i<2 else 9)
    setrow(table.rows[0], ['Programmer','Modules','System','Case\nManager','Agency\nFocal Person','Adminis-\ntrator','OFW'], True, True)
    table.rows[0]._tr.get_or_add_trPr().append(xml('tblHeader'))
    for m in modules:
        start=len(table.rows)
        setrow(table.add_row(), ['[Programmer]',m['name'],'','','','',''],True,True)
        for n,(text,roles) in enumerate(m['rows'],1):
            vals=['',f'{n}.  {text}']+[('*' if str(i) in roles else ('!' if (i==4 and '!' in roles) or (i==2 and '?' in roles) else '')) for i in range(5)]
            setrow(table.add_row(),vals,keep=(n==len(m['rows'])))
        # Reference-style programmer field spanning module heading and functions.
        first=table.cell(start,0)
        first.merge(table.cell(len(table.rows)-1,0))
        # merge leaves empty paragraphs; keep the placeholder only.
        first.text='[Programmer]'
        first.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
        for r in first.paragraphs[0].runs: r.font.size=Pt(10)
        row=table.add_row()
        setrow(row,['No. of Points','','','','','',''])
        row.cells[0].merge(row.cells[1]); row.cells[2].merge(row.cells[6])
        row.cells[0].text='No. of Points'; row.cells[2].text=''
    row=table.add_row()
    setrow(row,['Total Number of Modules','',str(len(modules)),'','','',''],True)
    row.cells[0].merge(row.cells[1]); row.cells[2].merge(row.cells[6])
    row.cells[0].text='Total Number of Modules'; row.cells[2].text=str(len(modules))
    row.cells[2].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    p=doc.add_paragraph('Access Notes'); p.runs[0].bold=True
    p.paragraph_format.space_after=Pt(8)
    for i,note in enumerate(notes,1):
        p=doc.add_paragraph(f'{i}. {note}')
        p.paragraph_format.space_after=Pt(7)
        for r in p.runs:r.font.size=Pt(10)
    doc.core_properties.title='Appendix N List of Modules'
    doc.core_properties.author=''
    doc.core_properties.subject='Implemented system module and actor matrix'
    target=OUT/'Appendix_N_List_of_Modules.docx'
    doc.save(target)
    (OUT/'appendix_n_matrix.json').write_text(json.dumps({'modules':modules,'notes':notes},indent=2),encoding='utf-8')
    print(f'{target}\n{len(modules)} modules, {sum(len(m["rows"]) for m in modules)} function rows')

if __name__=='__main__': build()
